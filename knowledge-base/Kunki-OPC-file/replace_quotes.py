#!/usr/bin/env python3
"""
替换docx文件中所有直引号 " 为标准中文弯引号 “”
保留原有格式不变，只替换引号字符
"""

from docx import Document
import re

def replace_quotes_in_docx(input_path, output_path):
    """
    替换docx文件中的直引号为中文弯引号
    """
    doc = Document(input_path)

    # 统计替换次数
    replace_count = 0

    # 处理所有段落
    for para in doc.paragraphs:
        if '"' in para.text:
            # 我们需要逐个替换，成对处理
            # 策略：奇数位" → “（左引号），偶数位" → ”（右引号）
            new_text = ''
            quote_count = 0
            for char in para.text:
                if char == '"':
                    quote_count += 1
                    if quote_count % 2 == 1:
                        new_text += '“'
                    else:
                        new_text += '”'
                else:
                    new_text += char
            replace_count += para.text.count('"')
            para.text = new_text

    # 处理表格中的内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '"' in para.text:
                        new_text = ''
                        quote_count = 0
                        for char in para.text:
                            if char == '"':
                                quote_count += 1
                                if quote_count % 2 == 1:
                                    new_text += '“'
                                else:
                                    new_text += '”'
                            else:
                                new_text += char
                        replace_count += para.text.count('"')
                        para.text = new_text

    # 保存文件
    doc.save(output_path)
    print(f"完成！共替换 {replace_count} 个直引号为中文弯引号")
    print(f"原始文件: {input_path}")
    print(f"输出文件: {output_path}")
    return replace_count

if __name__ == "__main__":
    input_file = "./conversations/conv-1777644786646/OpenClaw一人公司 AI自媒体运营实战手册.docx"
    output_file = "./conversations/conv-1777644786646/OpenClaw一人公司 AI自媒体运营实战手册_引号修正版.docx"

    replace_quotes_in_docx(input_file, output_file)

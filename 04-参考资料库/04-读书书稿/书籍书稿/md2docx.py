#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt
import re

def md_to_docx(md_file_path, docx_file_path):
    # 读取markdown内容
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 创建docx文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)

    # 按行处理markdown
    lines = md_content.split('\n')
    in_code_block = False
    in_table = False

    for line in lines:
        line = line.rstrip()

        # 处理代码块标记
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            # 代码块内容保持原样
            p = doc.add_paragraph('    ' + line)
            p.style = 'No Spacing'
            continue

        # 处理标题
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            continue

        # 处理分隔线
        if line.strip() == '---':
            doc.add_paragraph('')
            continue

        # 处理表格边框线（|------|...|），跳过
        if re.match(r'^\|[\s\-]+\|', line.strip()):
            continue

        # 处理表格行
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # 简单处理，转换为段落
            content = line.strip().strip('|')
            cells = [cell.strip() for cell in content.split('|')]
            doc.add_paragraph(' | '.join(cells))
            continue

        # 处理无序列表
        stripped = line.strip()
        if stripped.startswith('- ') or stripped.startswith('* '):
            content = stripped[2:]
            # 处理格式
            content = clean_markdown(content)
            doc.add_paragraph(content, style='List Bullet')
            continue

        # 处理有序列表
        match = re.match(r'^\d+\.\s+', stripped)
        if match:
            content = stripped[match.end():]
            content = clean_markdown(content)
            doc.add_paragraph(content, style='List Number')
            continue

        # 处理空行
        if not stripped:
            continue

        # 普通段落
        content = clean_markdown(line)
        doc.add_paragraph(content)

    # 保存文档
    doc.save(docx_file_path)
    print(f"转换完成：{docx_file_path}")

def clean_markdown(content):
    # 移除markdown格式符号（简化处理）
    # 移除加粗标记 ** **
    content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
    # 移除斜体标记 * *
    content = re.sub(r'\*(.*?)\*', r'\1', content)
    # 移除链接标记 [text](url) → text
    content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', content)
    # 移除反引号 `code` → code
    content = re.sub(r'`(.*?)`', r'\1', content)
    return content

if __name__ == "__main__":
    input_md = "OpenClaw一人公司 AI自媒体运营实战手册.md"
    output_docx = "OpenClaw一人公司 AI自媒体运营实战手册.docx"
    md_to_docx(input_md, output_docx)
